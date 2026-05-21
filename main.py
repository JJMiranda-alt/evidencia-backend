"""
EvidencIA — Backend FastAPI
Sistema de Gestión de Evaluaciones EC0301 CONOCER
Developed by Miranda
Versión 2.1 — Evaluación completa: E1 + E2 + E3 + Conocimientos
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import anthropic
import docx
import io
import os
import json
from datetime import datetime

app = FastAPI(
    title="EvidencIA API",
    description="Sistema de Gestión de Evaluaciones CONOCER EC0301 — Developed by Miranda",
    version="2.1.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

# ══════════════════════════════════════════════
# PROMPTS POR ELEMENTO
# ══════════════════════════════════════════════

PROMPT_E1 = """Eres evaluador experto del CONOCER para el estándar EC0301 "Diseño de cursos de formación del capital humano".
Evalúa la Carta Descriptiva (Elemento 1) con los 48 reactivos del IEC.

REACTIVOS E1 (Lista de Cotejo 1):
P1-Presentación(16): 1.1 Formato digital, 2.2 Nombre del curso, 3.3 Nombre instructor, 4.4 Fecha y horario congruentes con 180min, 5.5 Perfil participante, 6.6 Número participantes, 7.7 Objetivos presentes, 8.8 4 momentos didácticos 180min, 9.9 Contenido temático, 10.10 Técnicas instruccionales diferenciadas de grupales, 11.11 3 técnicas grupales(rompehielo Encuadre+energizante Desarrollo+cierre grupal Cierre) en columna grupal, 12.12 Actividades instructor/participante, 13.13 Estrategias evaluativas 3 momentos, 14.14 Materiales por etapa, 15.15 Tiempos=180min exactos, 16.16 Sin errores ortográficos
P2-Obj.General(4×4.16pts): 17.1 Sujeto, 18.2 Conducta observable, 19.3 Condición, 20.4 Criterio cuantificable(% o número concreto)
P3-Obj.Específicos(6×4.16pts): 21.1 Sujeto todos, 22.2 Conducta todos, 23.3 Condición todos, 24.4 Criterio cuantificable todos, 25.5 Congruentes contenido, 26.6 Congruentes perfil
P4-Contenido(3×0.58): 27.1 Congruente propósito, 28.2 Alineado objetivos, 29.3 Secuencia progresiva
P5-Téc.Instruc.(3×0.58): 30.1 Pertinentes objetivos, 31.2 Adecuadas perfil, 32.3 Pertinentes número
P6-Téc.Grupales(3×0.58): 33.1 Congruentes etapa/propósito, 34.2 Adecuadas perfil, 35.3 Pertinentes número
P7-Actividades(4×0.58): 36.1 Diferenciadas por etapa, 37.2 Congruentes contenido, 38.3 Acordes perfil, 39.4 Descripción clara y tiempos
P8-Evaluación(4×0.58+1×4.16): 40.1 Alineadas momentos, 41.2 Criterios definidos, 42.3 3 momentos presentes, 43.4 Instrumentos por etapa, 44.5 Evidencias sumativas con características y criterio aceptación(4.16pts)
P9-Materiales(3×0.58): 45.1 Corresponden actividades, 46.2 Adecuados perfil, 47.3 Pertinentes número
AHV(0/-0.58): 48.1 Secuencia progresiva simple→complejo (0 si SI, -0.58 si NO)

Pesos: P1 reactivos 1.1,2.2,3.3,4.4,6.6,7.7,16.16=0.25; 5.5,8.8-15.15=0.58 excepto 13.13=4.16; P2,P3=4.16c/u; P4-P9=0.58 excepto 44.5=4.16

Responde SOLO con JSON:
{
  "elemento": "E1",
  "candidato": "nombre del instructor",
  "curso": "nombre del curso",
  "fecha": "fecha y horario",
  "pts_total": 00.00,
  "pct": 00.0,
  "max_pts": 61.03,
  "reactivos": {"1.1":{"resultado":"SI","obs":"breve"},...todos 48},
  "pendientes": ["reactivo — descripción corrección"],
  "retroalimentacion": "fortalezas, áreas de oportunidad y recomendaciones detalladas",
  "dictamen": "COMPETENTE" o "AÚN NO COMPETENTE"
}"""

PROMPT_E2 = """Eres evaluador experto del CONOCER para EC0301. Evalúa los Instrumentos de Evaluación (Elemento 2) con los 31 reactivos.

REACTIVOS E2 (Lista de Cotejo 2):
P1-Instrumentos elaborados(10): 49.1 Nombre del curso(0.25), 50.2 Espacio nombre instructor(0.25), 51.3 Espacio nombre participante(0.25), 52.4 Espacio fecha aplicación(0.25), 53.5 Instrucciones de aplicación(0.58), 54.6 Reactivos de evaluación(0.58), 55.7 Claves de respuestas evaluador(0.58), 56.8 Corresponden estrategias carta descriptiva(0.25), 57.9 Formato digital/impreso(0.25), 58.10 Sin errores ortográficos(0.25)
P2-Instrucciones aplicación(4): 59.1 Condiciones de aplicación(0.25), 60.2 Tiempos evaluación(0.25), 61.3 Indicaciones participante(0.58), 62.4 Indicaciones evaluador(0.58)
P3-Reactivos instrumento(5): 63.1 Corresponden objetivos aprendizaje(4.16), 64.2 Congruentes tipo instrumento(0.58), 65.3 Verifican una sola evidencia/característica(0.58), 66.4 Son medibles(0.58), 67.5 Indican su valor(0.25)
P4-Claves respuestas(3): 68.1 Respuestas definidas correctas(0.58), 69.2 Ponderación cada reactivo(0.25), 70.3 Puntaje total esperado(0.25)
P5-Instrumento satisfacción(9): 71.1 Nombre del curso(0.25), 72.2 Nombre instructor(0.25), 73.3 Instrucciones generales(0.58), 74.4 Escala estimación satisfacción(0.25), 75.5 Reactivos características evento(0.25), 76.6 Reactivos contenido curso(0.25), 77.7 Reactivos materiales didácticos(0.25), 78.8 Reactivos desempeño instructor(0.25), 79.9 Espacios comentarios(0.25)

Responde SOLO con JSON:
{
  "elemento": "E2",
  "candidato": "nombre",
  "curso": "nombre curso",
  "pts_total": 00.00,
  "pct": 00.0,
  "max_pts": 34.46,
  "reactivos": {"49.1":{"resultado":"SI","obs":"breve"},...todos 31},
  "pendientes": ["reactivo — corrección"],
  "retroalimentacion": "fortalezas, áreas de oportunidad y recomendaciones",
  "dictamen": "COMPETENTE" o "AÚN NO COMPETENTE"
}"""

PROMPT_E3 = """Eres evaluador experto del CONOCER para EC0301. Evalúa los Manuales del Curso (Elemento 3) con los 60 reactivos.

REACTIVOS E3 (Lista de Cotejo 3):
P1-Manual participante(11×0.25 excepto 85.6=0.58): 80.1 Nombre curso, 81.2 Nombre diseñador, 82.3 Índice, 83.4 Presentación, 84.5 Introducción, 85.6 Objetivo general acorde carta(0.58), 86.7 Objetivos específicos acordes carta, 87.8 Temas desglosados, 88.9 Fuentes información, 89.10 Formato digital/impreso, 90.11 Sin errores ortográficos
P2-Presentación manual(3): 91.1 Bienvenida participante(0.25), 92.2 Recomendaciones uso manual(0.25), 93.3 Organización del manual(0.58)
P3-Introducción manual(4): 94.1 Resumen temas(0.58), 95.2 Beneficio curso participantes(0.58), 96.3 Enfoque didáctico(0.25), 97.4 Congruente objetivo aprendizaje(0.25)
P4-Temas desarrollados(7×0.58): 98.1 Corresponden carta descriptiva, 99.2 Congruentes objetivos, 100.3 Mencionan objetivos específicos, 101.4 De simple a complejo, 102.5 Actividades desarrollo tema, 103.6 Síntesis/conclusiones, 104.7 Forma evaluación por tema
P5-Fuentes participante(6×0.25): 105.1 Corresponden objetivos, 106.2 Nombre autor, 107.3 Año publicación/fecha acceso, 108.4 Título obra, 109.5 Editorial/URL, 110.6 País origen
P6-Manual instructor(12): 111.1-112.2 Nombre curso/diseñador(0.25c/u), 113.3 Índice(0.25), 114.4 Introducción(0.58), 115.5 Carta descriptiva(0.25), 116.6 Requerimientos lugar(0.25), 117.7 Sugerencias temas(0.58), 118.8 Instrumentos evaluación(0.58), 119.9 Claves respuestas(0.58), 120.10 Fuentes info(0.25), 121.11 Formato digital/impreso(0.25), 122.12 Sin errores ortográficos(0.25)
P7-Introducción instructor(3): 123.1 Propósito manual(0.58), 124.2 Estructura curso(0.25), 125.3 Modalidad curso(0.25)
P8-Requerimientos lugar(4): 126.1 Características lugar(0.58), 127.2 Material apoyo(0.25), 128.3 Equipo necesario(0.25), 129.4 Recomendaciones uso material(0.25)
P9-Temas instructor(4×0.58): 130.1 Corresponden carta, 131.2 Sugerencias apoyos, 132.3 Técnicas/actividades/ejemplos, 133.4 Formas/criterios/tiempos evaluación
P10-Fuentes instructor(6×0.25): 134.1 Corresponden objetivos, 135.2 Nombre autor, 136.3 Año/fecha acceso, 137.4 Título obra, 138.5 Editorial/URL, 139.6 País origen

Responde SOLO con JSON:
{
  "elemento": "E3",
  "candidato": "nombre",
  "curso": "nombre curso",
  "pts_total": 00.00,
  "pct": 00.0,
  "max_pts": 1.93,
  "reactivos": {"80.1":{"resultado":"SI","obs":"breve"},...todos 60},
  "pendientes": ["reactivo — corrección"],
  "retroalimentacion": "fortalezas, áreas de oportunidad y recomendaciones",
  "dictamen": "COMPETENTE" o "AÚN NO COMPETENTE"
}"""

PROMPT_CONOCIMIENTOS = """Eres evaluador experto del CONOCER para EC0301. Evalúa el Examen de Conocimientos con los 6 reactivos de los 2 cuestionarios.

CUESTIONARIO 1 — Diseño del Curso (4 reactivos × 0.25 pts = 1.00 pt):
Reactivo 140.1 — Teorías del aprendizaje (4 preguntas): Conductismo=cambio conducta por repetición, Cognitivismo=ente activo procesa información, Constructivismo=reestructura conocimientos previos, Humanismo=deseo natural+afectividad+intelecto
Reactivo 141.1 — Principios educación adultos (6 preguntas): Necesidad de saber=identificar estado conocimiento, Desaprendizaje=modelos mejores de desempeño, Motivación=reforzamientos positivos sin sobrevaloraciones, Recuperación experiencia=ejemplos de experiencias anteriores, Desaprendizaje=examinar hábitos y prejuicios, Aplicación práctica=contenidos en contexto cercano realidad
Reactivo 142.1 — Técnicas instruccionales (6 preguntas): Expositiva=un solo sentido cuidadosa preparación, Expositiva=disertación formal directa, Diálogo/Discusión=intercambio libre verbal, Diálogo/Discusión=compartir ideas identificar problemas, Demostrativa=observar corregir errores realimentar, Demostrativa=habilidades manuales
Reactivo 143.1 — Técnicas grupales (3 preguntas): Rompehielo=clima confianza comprensión, Energizante=cansancio monotonía, Cierre=reflexionar concluir

CUESTIONARIO 2 — Instrumentos de Evaluación (2 reactivos × 0.58 pts = 1.16 pts):
Reactivo 144.1 — Validez y confiabilidad V/F: a)V-confiabilidad=resultados similares poblaciones similares, b)V-validez=evalúa todos los aspectos, c)F-validez NO depende de algunos aspectos, d)F-confiabilidad NO es en poblaciones distintas
Reactivo 145.1 — Tipos instrumentos V/F: a)V-desempeños inicia verbo presente, b)F-productos NO más de una característica, c)V-lista cotejo una sola característica, d)V-cotejo registrar como sí, e)F-cuestionario NO más de una respuesta por reactivo, f)V-escolaridad candidatos, g)V-clave respuestas para todo reactivo, h)F-NO se recomiendan negaciones

Evalúa las respuestas del candidato y determina SI/NO por reactivo completo (todos correctos=SI).

Responde SOLO con JSON:
{
  "elemento": "CONOCIMIENTOS",
  "candidato": "nombre",
  "pts_total": 0.00,
  "pct": 0.0,
  "max_pts": 2.66,
  "reactivos": {
    "140.1":{"resultado":"SI","obs":"todas correctas"},
    "141.1":{"resultado":"NO","obs":"falló pregunta X"},
    "142.1":{"resultado":"SI","obs":""},
    "143.1":{"resultado":"SI","obs":""},
    "144.1":{"resultado":"SI","obs":""},
    "145.1":{"resultado":"NO","obs":"inciso b incorrecto"}
  },
  "pendientes": ["reactivo — temas a repasar"],
  "retroalimentacion": "resultado y recomendaciones de estudio",
  "dictamen": "APROBADO" o "REPROBADO"
}"""

PROMPTS = {"E1": PROMPT_E1, "E2": PROMPT_E2, "E3": PROMPT_E3, "CONOCIMIENTOS": PROMPT_CONOCIMIENTOS}

PESOS_E1 = {
    "1.1":0.58,"2.2":0.25,"3.3":0.25,"4.4":0.25,"5.5":0.58,"6.6":0.25,
    "7.7":0.58,"8.8":0.58,"9.9":0.58,"10.10":0.58,"11.11":0.58,"12.12":0.58,
    "13.13":4.16,"14.14":0.58,"15.15":0.58,"16.16":0.25,
    "17.1":4.16,"18.2":4.16,"19.3":4.16,"20.4":4.16,
    "21.1":4.16,"22.2":4.16,"23.3":4.16,"24.4":4.16,"25.5":4.16,"26.6":0.58,
    "27.1":0.58,"28.2":0.25,"29.3":0.25,
    "30.1":0.58,"31.2":0.25,"32.3":0.25,
    "33.1":0.25,"34.2":0.25,"35.3":0.25,
    "36.1":4.16,"37.2":0.58,"38.3":0.58,"39.4":0.58,
    "40.1":4.16,"41.2":0.58,"42.3":0.58,"43.4":0.58,"44.5":4.16,
    "45.1":0.25,"46.2":0.25,"47.3":0.25,"48.1":0,
}


def extraer_texto_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto = cell.text.strip()
                if texto and texto not in parrafos:
                    parrafos.append(texto)
    return "\n".join(parrafos)


def parsear_respuesta(texto: str) -> dict:
    texto = texto.strip()
    if "```" in texto:
        partes = texto.split("```")
        for p in partes:
            p = p.strip()
            if p.startswith("json"):
                p = p[4:].strip()
            if p.startswith("{"):
                texto = p
                break
    return json.loads(texto.strip())


@app.get("/")
async def root():
    return {"sistema":"EvidencIA API","version":"2.1.0","desarrollado_por":"Miranda",
            "estandar":"EC0301 CONOCER","elementos":["E1","E2","E3","CONOCIMIENTOS"],"status":"activo"}


@app.get("/health")
async def health():
    return {"status":"ok","timestamp":datetime.now().isoformat()}


@app.post("/evaluar")
async def evaluar(
    archivo: UploadFile = File(...),
    candidato_id: str = Form(default=""),
    candidato_nombre: str = Form(default=""),
    elemento: str = Form(default="E1"),
):
    elemento = elemento.upper()
    if elemento not in PROMPTS:
        raise HTTPException(400, f"Elemento no válido. Use: E1, E2, E3 o CONOCIMIENTOS")

    if not archivo.filename.endswith(('.docx', '.pdf')):
        raise HTTPException(400, "Solo se aceptan archivos .docx")

    contenido = await archivo.read()
    if len(contenido) > 10 * 1024 * 1024:
        raise HTTPException(400, "Archivo demasiado grande (máximo 10MB)")

    try:
        texto = extraer_texto_docx(contenido)
    except Exception as e:
        raise HTTPException(422, f"Error al leer el archivo: {str(e)}")

    if len(texto.strip()) < 50:
        raise HTTPException(422, "El archivo parece estar vacío")

    try:
        mensaje = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            system=PROMPTS[elemento],
            messages=[{
                "role": "user",
                "content": f"Candidato: {candidato_nombre or 'No especificado'}\nElemento a evaluar: {elemento}\n\nCONTENIDO DEL DOCUMENTO:\n{texto}\n\nResponde SOLO con el JSON de evaluación."
            }]
        )

        resultado = parsear_respuesta(mensaje.content[0].text)
        resultado["candidato_id"] = candidato_id
        resultado["fecha_evaluacion"] = datetime.now().strftime("%d/%m/%Y %H:%M")
        resultado["archivo"] = archivo.filename
        resultado["evaluador"] = "Dr. José Juan Miranda Torres"

        return JSONResponse(content=resultado)

    except json.JSONDecodeError as e:
        raise HTTPException(500, f"Error al parsear respuesta de IA: {str(e)}")
    except anthropic.AuthenticationError:
        raise HTTPException(401, "API Key de Anthropic inválida")
    except anthropic.RateLimitError:
        raise HTTPException(429, "Límite de requests alcanzado — espera un momento")
    except Exception as e:
        raise HTTPException(500, f"Error: {str(e)}")
